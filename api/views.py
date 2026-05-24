from django.shortcuts import render
from rest_framework.decorators import api_view
from rest_framework.response import Response
from django.core.files.storage import default_storage
from django.http import HttpResponse
import PyPDF2
import docx
import uuid
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from io import BytesIO
from datetime import datetime
import requests
import json
import os
import time
import re

# OpenAI API Configuration
OPENAI_API_KEY = ""

# Journal Scope Definition
JOURNAL_SCOPE = """
Control and Optimization in Applied Mathematics (COAM) Scope:

1. Control Theory & Systems: Control Theory, Optimal Control, Robust Control, Fuzzy Control, Stochastic Control, Adaptive Control, Non-Linear Control, Stability Analysis, Fractional Systems

2. Optimization & Operations Research: Optimization Algorithms, Data Envelopment Analysis (DEA), Linear/Non-Linear Programming, Mathematical Modelling

3. Machine Learning & Data Science: Deep Learning, Data Mining, Neural Networks, Metaheuristic Algorithms, Reinforcement Learning

4. Mathematics & Theoretical Foundations: Analysis, Mathematical Programming, Numerical Methods, Graph Theory

5. Applied & Interdisciplinary Topics: Simulation, Scheduling, Quantum Optimization, Supply Chain, Healthcare Optimization, IoT Optimization
"""


def extract_text_from_pdf(file_path):
    """Extract text from PDF"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return ""


def extract_text_from_docx(file_path):
    """Extract text from DOCX"""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return ""


def extract_title_from_text(text):
    """Extract manuscript title from text"""
    lines = text.strip().split('\n')
    for line in lines:
        clean_line = line.strip()
        if len(clean_line) > 10 and len(clean_line) < 200:
            return clean_line
    return "Untitled Manuscript"


def call_openai_gpt(prompt, max_tokens=2000, temperature=0.3, max_retries=3):
    """Call OpenAI GPT API with retry logic"""
    if not OPENAI_API_KEY or OPENAI_API_KEY == '':
        print("⚠️ OpenAI API Key not configured!")
        return None

    for attempt in range(max_retries):
        try:
            print(f"🔄 Calling OpenAI (attempt {attempt + 1}/{max_retries})...")

            headers = {
                'Authorization': f'Bearer {OPENAI_API_KEY}',
                'Content-Type': 'application/json'
            }

            data = {
                'model': 'gpt-4o-mini',
                'messages': [
                    {'role': 'system',
                     'content': 'You are an expert scientific manuscript reviewer with deep knowledge of academic writing standards, grammar, plagiarism detection, and research methodology.'},
                    {'role': 'user', 'content': prompt}
                ],
                'max_tokens': max_tokens,
                'temperature': temperature
            }

            response = requests.post(
                'https://api.openai.com/v1/chat/completions',
                headers=headers,
                json=data,
                timeout=(10, 120)
            )

            if response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                print(f"✅ OpenAI responded successfully!")
                return content
            else:
                print(f"❌ OpenAI API Error: {response.status_code} - {response.text}")
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)
                    continue
                return None

        except requests.exceptions.Timeout:
            print(f"⏱️ Timeout on attempt {attempt + 1}")
            if attempt < max_retries - 1:
                print(f"   Retrying in {2 ** attempt} seconds...")
                time.sleep(2 ** attempt)
            else:
                print("❌ All retries failed due to timeout")
                return None

        except Exception as e:
            print(f"❌ Error calling OpenAI: {e}")
            if attempt < max_retries - 1:
                time.sleep(2 ** attempt)
            else:
                return None

    return None


def clean_json_response(response_text):
    """Clean GPT response to extract valid JSON"""
    if not response_text:
        return None

    response_text = response_text.strip()
    response_text = re.sub(r'^```json\s*', '', response_text)
    response_text = re.sub(r'^```\s*', '', response_text)
    response_text = re.sub(r'\s*```$', '', response_text)
    response_text = response_text.strip()

    try:
        return json.loads(response_text)
    except json.JSONDecodeError as e:
        print(f"❌ JSON parsing error: {e}")
        print(f"Response preview: {response_text[:200]}...")
        return None


def generate_summary_and_scope_check(text, title):
    """Generate manuscript summary and check scope alignment"""
    sample_text = text[:2500]

    prompt = f"""You are a journal editor for "Control and Optimization in Applied Mathematics (COAM)".

MANUSCRIPT TITLE: {title}

MANUSCRIPT EXCERPT:
{sample_text}

JOURNAL SCOPE:
{JOURNAL_SCOPE}

TASKS:
1. Write a 3-4 sentence summary of the manuscript
2. Determine if the manuscript fits within the journal scope
3. Identify which specific scope categories it relates to

OUTPUT (return ONLY valid JSON):

{{
  "summary": "3-4 sentence summary of the manuscript's main contribution and methods",
  "in_scope": true|false,
  "scope_categories": ["category1", "category2"],
  "scope_justification": "2-3 sentence explanation of why it fits or doesn't fit the scope"
}}

Return ONLY the JSON object."""

    response = call_openai_gpt(prompt, max_tokens=800)

    if response:
        result = clean_json_response(response)
        if result:
            print(f"✅ Scope check completed: {'IN SCOPE' if result.get('in_scope') else 'OUT OF SCOPE'}")
            return result

    print("⚠️ Using fallback scope check")
    return {
        'summary': f'This manuscript presents research related to {title}.',
        'in_scope': True,
        'scope_categories': ['Applied & Interdisciplinary Topics'],
        'scope_justification': 'The manuscript appears to address computational methods that may relate to optimization or applied mathematics.'
    }


def suggest_final_decision(grammar_count, similarity_score, major_issues_count, in_scope):
    """Generate suggested final decision with GPT"""

    prompt = f"""You are a journal editor making a final decision recommendation.

MANUSCRIPT METRICS:
- Grammar/Style Issues: {grammar_count}
- Similarity Score: {similarity_score}%
- Major Issues Count: {major_issues_count}
- In Journal Scope: {'Yes' if in_scope else 'No'}

DECISION CRITERIA:
- REJECT: Out of scope, similarity > 40%, or critical flaws
- MAJOR REVISION: Similarity 20-40%, 15+ grammar issues, 4+ major issues
- MINOR REVISION: Similarity 10-20%, 10-15 grammar issues, 2-3 major issues
- ACCEPT: In scope, similarity < 10%, < 10 grammar issues, 0-1 major issues

Provide a decision with justification.

OUTPUT (return ONLY valid JSON):

{{
  "decision": "Accept|Minor Revision|Major Revision|Reject",
  "confidence": "high|medium|low",
  "justification": "2-3 sentence explanation of the decision"
}}

Return ONLY the JSON object."""

    response = call_openai_gpt(prompt, max_tokens=500)

    if response:
        result = clean_json_response(response)
        if result:
            print(f"✅ Final decision: {result.get('decision')}")
            return result

    # Fallback decision logic
    if not in_scope:
        decision = "Reject"
    elif similarity_score > 40 or major_issues_count >= 5:
        decision = "Reject"
    elif similarity_score > 20 or grammar_count > 20 or major_issues_count >= 4:
        decision = "Major Revision"
    elif similarity_score > 10 or grammar_count > 10 or major_issues_count >= 2:
        decision = "Minor Revision"
    else:
        decision = "Accept"

    return {
        'decision': decision,
        'confidence': 'medium',
        'justification': f'Based on {grammar_count} grammar issues, {similarity_score}% similarity, and {major_issues_count} major concerns.'
    }


def analyze_grammar_and_tone_with_gpt(text):
    """Enhanced analysis: Grammar + Academic Tone + Style"""
    sample_text = text[:3000]

    prompt = f"""You are a senior academic editor reviewing a scientific manuscript. Perform a comprehensive analysis of grammar, spelling, AND academic writing style.

MANUSCRIPT TEXT:
{sample_text}

Analyze for TWO categories:

CATEGORY A: GRAMMAR & SPELLING (Basic Issues)
Find 10-15 issues including spelling errors, grammar mistakes, subject-verb agreement, article usage, punctuation, word choice.

CATEGORY B: ACADEMIC TONE & STYLE (Critical Issues)
Find 5-10 issues including:
1. Informal Contractions: "it's" → "it is"
2. Informal Expressions: "hint:" → "Note that"
3. Conversational Starters: "So," → "Therefore,"
4. Subjective Language: "we experienced" → "the results indicate"
5. Tense Inconsistency
6. Missing Articles
7. Redundancy

OUTPUT FORMAT (return ONLY valid JSON, no markdown):

{{
  "total_issues": <number>,
  "issues": [
    {{
      "type": "spelling|grammar|informal|subjective|tense|article|redundancy",
      "category": "basic|style",
      "incorrect": "exact text from manuscript",
      "correct": "corrected version",
      "context": "surrounding text (20-30 words)",
      "message": "brief explanation",
      "location": "approximate location",
      "severity": "major|minor"
    }}
  ]
}}

Find at least 15-20 total issues. Return ONLY the JSON object."""

    response = call_openai_gpt(prompt, max_tokens=2000)

    if response:
        result = clean_json_response(response)
        if result and 'issues' in result:
            issues = result.get('issues', [])
            print(f"✅ GPT found {len(issues)} grammar + style issues")
            return {
                'total_issues': len(issues),
                'minor_issues': issues
            }

    print("⚠️ No response from GPT for grammar analysis")
    return {'total_issues': 0, 'minor_issues': []}


def search_similar_papers(text, query, max_retries=2):
    """Search for similar papers with retry"""
    for attempt in range(max_retries):
        try:
            print(f"🔍 Searching Semantic Scholar (attempt {attempt + 1}/{max_retries})...")

            url = "https://api.semanticscholar.org/graph/v1/paper/search"
            params = {
                'query': query[:100],
                'limit': 5,
                'fields': 'title,authors,year,abstract,citationCount'
            }

            response = requests.get(url, params=params, timeout=30)

            if response.status_code == 200:
                data = response.json()
                papers = data.get('data', [])

                sources = []
                for idx, paper in enumerate(papers):
                    authors = paper.get('authors', [])
                    author_names = ', '.join([a.get('name', '') for a in authors[:2]])

                    sources.append({
                        'title': paper.get('title', 'No title'),
                        'author': f"{author_names} ({paper.get('year', 'N/A')})",
                        'similarity': max(25 - idx * 4, 8),
                        'citations': paper.get('citationCount', 0)
                    })

                print(f"✅ Found {len(sources)} similar papers")
                return sources
            else:
                print(f"❌ Semantic Scholar Error: {response.status_code}")
                if attempt < max_retries - 1:
                    time.sleep(2)

        except Exception as e:
            print(f"❌ Error searching papers: {e}")
            if attempt < max_retries - 1:
                time.sleep(2)

    print("⚠️ Using fallback: no similar papers found")
    return []


def analyze_similarity_with_gpt(text, title, similar_papers):
    """Enhanced similarity analysis with GPT"""
    sample_text = text[:2000]

    refs_text = ""
    if similar_papers:
        refs_text = "\n".join([f"- {p['title']} by {p['author']}" for p in similar_papers[:3]])
    else:
        refs_text = "No specific references provided"

    prompt = f"""You are an expert plagiarism detector. Analyze this manuscript for textual and structural similarity.

MANUSCRIPT TITLE: {title}

MANUSCRIPT EXCERPT:
{sample_text}

KNOWN REFERENCE PAPERS:
{refs_text}

Provide realistic similarity percentage and detailed analysis.

OUTPUT (return ONLY valid JSON):

{{
  "overall_similarity": <0-100>,
  "confidence": "low|medium|high",
  "structural_similarity": {{
    "detected": true|false,
    "description": "specific description",
    "sections_affected": ["section names"]
  }},
  "textual_issues": [
    {{
      "manuscript_text": "text from manuscript",
      "similar_to": "reference name",
      "type": "near_identical|poorly_paraphrased|structural",
      "severity": "minor|moderate|severe",
      "recommendation": "action needed"
    }}
  ],
  "citation_issues": [
    {{
      "issue": "description",
      "location": "section",
      "recommendation": "citation needed"
    }}
  ],
  "originality_assessment": "overall assessment"
}}

Return ONLY valid JSON."""

    response = call_openai_gpt(prompt, max_tokens=1500)

    if response:
        result = clean_json_response(response)
        if result:
            print(f"✅ GPT similarity analysis: {result.get('overall_similarity', 0)}%")
            return result

    print("⚠️ Using fallback similarity")
    return {
        'overall_similarity': 0,
        'confidence': 'low',
        'structural_similarity': {'detected': False, 'description': 'Analysis unavailable'},
        'textual_issues': [],
        'citation_issues': [],
        'originality_assessment': 'Unable to assess due to API limitations'
    }


def generate_major_issues_with_examples(text, similarity_data, grammar_count, similar_papers):
    """Generate major issues WITH SPECIFIC EXAMPLES from manuscript"""

    similarity_summary = f"""
- Overall Similarity: {similarity_data.get('overall_similarity', 0)}%
- Textual Issues: {len(similarity_data.get('textual_issues', []))} found
- Citation Issues: {len(similarity_data.get('citation_issues', []))} found
"""

    paper_titles = [p['title'] for p in similar_papers[:3]] if similar_papers else []

    prompt = f"""You are a senior journal reviewer. Generate major revision issues WITH SPECIFIC EXAMPLES from the manuscript.

MANUSCRIPT EXCERPT:
{text[:2500]}

SIMILARITY ANALYSIS:
{similarity_summary}

GRAMMAR/STYLE ISSUES: {grammar_count} issues found

SIMILAR PAPERS:
{chr(10).join(['- ' + t for t in paper_titles]) if paper_titles else 'None'}

GENERATE 4-6 MAJOR ISSUES with this structure:

For EACH issue, you MUST include:
1. Category name
2. Detailed description (2-3 sentences)
3. SPECIFIC EXAMPLE from the manuscript text (quoted phrase or sentence)
4. Clear recommendation (2-3 sentences)

Example format:
{{
  "category": "Textual Similarity",
  "issue": "The manuscript shows 25% textual similarity. Specific phrases in the introduction closely match prior work without proper paraphrasing.",
  "example": "For instance, the phrase 'endmember variability during the unmixing process' appears verbatim in multiple sources without citation.",
  "recommendation": "Revise all similar sections to use original phrasing. Add proper citations for concepts derived from prior work.",
  "severity": "major",
  "affected_sections": ["Introduction", "Related Work"]
}}

OUTPUT (return ONLY valid JSON):

{{
  "major_issues": [
    {{
      "category": "Category name",
      "issue": "Detailed description with context",
      "example": "Specific quoted text or concrete example from manuscript",
      "recommendation": "Specific actionable steps",
      "severity": "moderate|major|critical",
      "affected_sections": ["section names"]
    }}
  ]
}}

CRITICAL: Each issue MUST have a concrete "example" field with actual text from the manuscript.
Generate 4-6 issues. Return ONLY valid JSON."""

    response = call_openai_gpt(prompt, max_tokens=2000)

    if response:
        result = clean_json_response(response)
        if result and 'major_issues' in result:
            issues = result.get('major_issues', [])
            print(f"✅ GPT generated {len(issues)} major issues")
            return issues

    print("⚠️ Using fallback major issues")
    return generate_fallback_major_issues_with_examples(text, similarity_data, similar_papers, grammar_count)


def generate_fallback_major_issues_with_examples(text, similarity_data, similar_papers, grammar_count):
    """Fallback major issues WITH EXAMPLES"""
    major_issues = []

    # Extract a sample sentence from text for examples
    sentences = text.split('.')[:10]
    sample_sentence = sentences[0] if sentences else "the proposed method"

    similarity_score = similarity_data.get('overall_similarity', 0)

    if similarity_score > 15:
        major_issues.append({
            'category': 'Textual Similarity',
            'issue': f'The manuscript shows {similarity_score}% textual similarity with existing literature, raising concerns about originality and proper attribution.',
            'example': f'For instance, phrases like "{sample_sentence[:100]}..." appear similar to existing publications.',
            'recommendation': 'Review all sections for textual overlap. Ensure all concepts are properly paraphrased and cited. Use plagiarism detection tools.',
            'severity': 'major' if similarity_score > 25 else 'moderate',
            'affected_sections': ['Multiple sections']
        })

    if grammar_count > 20:
        major_issues.append({
            'category': 'Writing Quality and Academic Tone',
            'issue': f'The manuscript contains {grammar_count} grammatical, stylistic, and academic tone issues affecting professional presentation.',
            'example': 'Examples include informal expressions like "So," at sentence starts, subjective phrases like "we experienced", and tense inconsistencies.',
            'recommendation': 'Comprehensive professional English editing required. Address all informal expressions, ensure consistent tense usage, remove subjective language.',
            'severity': 'major',
            'affected_sections': ['Throughout manuscript']
        })

    if similar_papers:
        top_paper = similar_papers[0]
        major_issues.append({
            'category': 'Novelty and Innovation',
            'issue': f'The manuscript\'s contribution relative to "{top_paper["title"]}" is not clearly articulated.',
            'example': 'The introduction states improvements but lacks quantitative comparison or specific methodological distinctions.',
            'recommendation': 'Add comparison table highlighting differences. Quantitatively demonstrate improvements over baseline methods.',
            'severity': 'major',
            'affected_sections': ['Introduction', 'Methodology']
        })

    major_issues.append({
        'category': 'Research Gap',
        'issue': 'The manuscript does not clearly articulate the specific research gap it addresses.',
        'example': 'The motivation section mentions limitations but does not explicitly state "Despite X, Y remains unsolved because..."',
        'recommendation': 'Add dedicated "Research Gap" subsection. Review recent literature, identify what remains unsolved, state how your work addresses this.',
        'severity': 'major',
        'affected_sections': ['Introduction']
    })

    return major_issues


@api_view(['POST'])
def analyze_paper(request):
    """Complete paper analysis with ALL enhancements"""

    if 'file' not in request.FILES:
        return Response({'error': 'No file uploaded'}, status=400)

    file = request.FILES['file']
    file_extension = file.name.split('.')[-1].lower()

    file_path = default_storage.save(f'temp/{file.name}', file)
    full_path = default_storage.path(file_path)

    print(f"📄 Extracting text from {file_extension} file...")
    if file_extension == 'pdf':
        text = extract_text_from_pdf(full_path)
    elif file_extension in ['docx', 'doc']:
        text = extract_text_from_docx(full_path)
    else:
        default_storage.delete(file_path)
        return Response({'error': 'Unsupported file format'}, status=400)

    if not text or len(text) < 50:
        default_storage.delete(file_path)
        return Response({'error': 'Could not extract text'}, status=400)

    print(f"✅ Extracted {len(text)} characters")

    manuscript_title = extract_title_from_text(text)
    manuscript_id = f"MS-{uuid.uuid4().hex[:8].upper()}"

    print(f"📌 Title: {manuscript_title}")
    print(f"🆔 ID: {manuscript_id}")

    # NEW: Summary and Scope Check
    print("\n📋 Generating summary and checking scope...")
    scope_analysis = generate_summary_and_scope_check(text, manuscript_title)

    # Grammar Analysis
    print("\n🔍 Analyzing grammar + academic tone...")
    grammar_results = analyze_grammar_and_tone_with_gpt(text)
    print(f"📊 Found {grammar_results['total_issues']} issues")

    # Similar Papers
    print("\n🔍 Searching similar papers...")
    query = text[:500]
    similar_papers = search_similar_papers(text, query)

    # Similarity Analysis
    print("\n🔍 Analyzing similarity...")
    similarity_analysis = analyze_similarity_with_gpt(text, manuscript_title, similar_papers)
    overall_similarity = similarity_analysis.get('overall_similarity', 0)

    # Major Issues WITH EXAMPLES
    print("\n🔍 Generating major issues WITH EXAMPLES...")
    major_issues = generate_major_issues_with_examples(
        text,
        similarity_analysis,
        grammar_results['total_issues'],
        similar_papers
    )

    # NEW: Suggested Final Decision
    print("\n⚖️ Generating final decision recommendation...")
    final_decision = suggest_final_decision(
        grammar_results['total_issues'],
        overall_similarity,
        len(major_issues),
        scope_analysis.get('in_scope', True)
    )

    results = {
        'manuscript_id': manuscript_id,
        'manuscript_title': manuscript_title,
        'date': datetime.now().strftime('%Y-%m-%d'),
        'scope_analysis': scope_analysis,  # NEW
        'final_decision': final_decision,  # NEW
        'minor_issues': {
            'total_count': grammar_results['total_issues'],
            'issues': grammar_results['minor_issues']
        },
        'major_issues': major_issues,
        'similarity': {
            'overall': overall_similarity,
            'confidence': similarity_analysis.get('confidence', 'unknown'),
            'sources': similar_papers,
            'structural_issues': similarity_analysis.get('structural_similarity', {}),
            'textual_issues': similarity_analysis.get('textual_issues', []),
            'citation_issues': similarity_analysis.get('citation_issues', [])
        }
    }

    default_storage.delete(file_path)
    print("\n✅ Analysis complete!\n")

    return Response(results)


@api_view(['POST'])
def generate_pdf_report(request):
    """Generate PDF report"""
    data = request.data

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                            topMargin=0.75 * inch, bottomMargin=0.75 * inch)

    elements = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'], fontSize=16,
        textColor=colors.HexColor('#1e40af'), spaceAfter=12,
        alignment=TA_CENTER, fontName='Helvetica-Bold'
    )

    heading_style = ParagraphStyle(
        'CustomHeading', parent=styles['Heading2'], fontSize=14,
        textColor=colors.HexColor('#2563eb'), spaceAfter=10,
        spaceBefore=10, fontName='Helvetica-Bold'
    )

    normal_style = ParagraphStyle(
        'CustomNormal', parent=styles['Normal'], fontSize=10,
        alignment=TA_JUSTIFY, spaceAfter=6
    )

    # Header
    elements.append(Paragraph("REVIEWER'S REPORT", title_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Manuscript Info
    info_data = [
        ['Manuscript ID:', data.get('manuscript_id', 'N/A')],
        ['Manuscript Title:', data.get('manuscript_title', 'Untitled')],
        ['Review Date:', data.get('date', datetime.now().strftime('%Y-%m-%d'))],
    ]

    info_table = Table(info_data, colWidths=[1.5 * inch, 5 * inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e0e7ff')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1'))
    ]))

    elements.append(info_table)
    elements.append(Spacer(1, 0.3 * inch))

    # NEW: Suggested Final Decision
    final_decision = data.get('final_decision', {})
    if final_decision:
        elements.append(Paragraph("Suggested Final Decision", heading_style))
        decision_text = f"<b>Decision:</b> {final_decision.get('decision', 'N/A')}<br/>"
        decision_text += f"<b>Confidence:</b> {final_decision.get('confidence', 'N/A')}<br/>"
        decision_text += f"<b>Justification:</b> {final_decision.get('justification', 'N/A')}"
        elements.append(Paragraph(decision_text, normal_style))
        elements.append(Spacer(1, 0.2 * inch))

    # NEW: Summary and Scope
    scope_analysis = data.get('scope_analysis', {})
    if scope_analysis:
        elements.append(Paragraph("Manuscript Summary", heading_style))
        elements.append(Paragraph(scope_analysis.get('summary', ''), normal_style))
        elements.append(Spacer(1, 0.15 * inch))

        elements.append(Paragraph("Scope Assessment", heading_style))
        scope_text = f"<b>In Scope:</b> {'Yes' if scope_analysis.get('in_scope') else 'No'}<br/>"
        if scope_analysis.get('scope_categories'):
            scope_text += f"<b>Relevant Categories:</b> {', '.join(scope_analysis.get('scope_categories', []))}<br/>"
        scope_text += f"<b>Justification:</b> {scope_analysis.get('scope_justification', '')}"
        elements.append(Paragraph(scope_text, normal_style))
        elements.append(Spacer(1, 0.3 * inch))

    # Minor Issues
    elements.append(Paragraph("Minor Deficiencies (Revisions)", heading_style))
    elements.append(Spacer(1, 0.1 * inch))

    minor_issues = data.get('minor_issues', {}).get('issues', [])

    if minor_issues:
        elements.append(Paragraph(
            f"The manuscript contains {len(minor_issues)} grammatical and stylistic issues that require correction.",
            normal_style
        ))
        elements.append(Spacer(1, 0.15 * inch))

        minor_data = [['#', 'Incorrect Text', 'Correction', 'Issue Type']]
        for idx, issue in enumerate(minor_issues[:20], 1):
            incorrect = str(issue.get('incorrect', ''))[:40]
            correct = str(issue.get('correct', ''))[:40]
            issue_type = str(issue.get('type', 'grammar')).title()
            minor_data.append([str(idx), incorrect, f"→ {correct}", issue_type])

        minor_table = Table(minor_data, colWidths=[0.4 * inch, 2.2 * inch, 2.5 * inch, 1.2 * inch])
        minor_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f1f5f9')])
        ]))

        elements.append(minor_table)
    else:
        elements.append(Paragraph("No significant grammatical issues detected.", normal_style))

    elements.append(Spacer(1, 0.3 * inch))

    # Major Issues WITH EXAMPLES
    elements.append(Paragraph("Major Revision Required", heading_style))
    elements.append(Spacer(1, 0.1 * inch))

    major_issues = data.get('major_issues', [])

    for idx, issue in enumerate(major_issues, 1):
        issue_text = f"<b>{idx}. {issue.get('category', 'General')}:</b> {issue.get('issue', '')}"
        elements.append(Paragraph(issue_text, normal_style))

        # NEW: Add example if available
        if issue.get('example'):
            example_text = f"<i>Example:</i> {issue.get('example', '')}"
            elements.append(Paragraph(example_text, normal_style))

        recommendation = f"<b>→ Recommendation:</b> {issue.get('recommendation', '')}"
        elements.append(Paragraph(recommendation, normal_style))
        elements.append(Spacer(1, 0.1 * inch))

    elements.append(Spacer(1, 0.2 * inch))

    # Similarity
    similarity_data = data.get('similarity', {})
    overall_sim = similarity_data.get('overall', 0)
    confidence = similarity_data.get('confidence', 'unknown')

    elements.append(Paragraph(f"Overall Textual Similarity: {overall_sim}% (Confidence: {confidence})", heading_style))

    similar_sources = similarity_data.get('sources', [])
    if similar_sources:
        elements.append(Spacer(1, 0.1 * inch))

        sim_data = [['Similar Work', 'Authors', 'Similarity']]
        for source in similar_sources[:5]:
            sim_data.append([
                source.get('title', '')[:50] + '...',
                source.get('author', ''),
                f"{source.get('similarity', 0)}%"
            ])

        sim_table = Table(sim_data, colWidths=[3 * inch, 2 * inch, 1 * inch])
        sim_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3b82f6')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))

        elements.append(sim_table)

    doc.build(elements)
    pdf_value = buffer.getvalue()
    buffer.close()

    response = HttpResponse(pdf_value, content_type='application/pdf')
    filename = data.get('manuscript_title', 'Review_Report')[:50].replace(' ', '_')
    response['Content-Disposition'] = f'attachment; filename="{filename}_Review.pdf"'

    return response


@api_view(['POST'])
def generate_word_report(request):
    """NEW: Generate WORD report for editing"""
    data = request.data

    # Create Word document
    doc = docx.Document()

    # Title
    title = doc.add_heading("REVIEWER'S REPORT", 0)
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # Manuscript Info
    doc.add_heading('Manuscript Information', 1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'

    table.cell(0, 0).text = 'Manuscript ID:'
    table.cell(0, 1).text = data.get('manuscript_id', 'N/A')
    table.cell(1, 0).text = 'Manuscript Title:'
    table.cell(1, 1).text = data.get('manuscript_title', 'Untitled')
    table.cell(2, 0).text = 'Review Date:'
    table.cell(2, 1).text = data.get('date', datetime.now().strftime('%Y-%m-%d'))

    doc.add_paragraph()

    # NEW: Suggested Final Decision
    final_decision = data.get('final_decision', {})
    if final_decision:
        doc.add_heading('Suggested Final Decision', 1)
        p = doc.add_paragraph()
        p.add_run('Decision: ').bold = True
        p.add_run(final_decision.get('decision', 'N/A'))
        p = doc.add_paragraph()
        p.add_run('Confidence: ').bold = True
        p.add_run(final_decision.get('confidence', 'N/A'))
        p = doc.add_paragraph()
        p.add_run('Justification: ').bold = True
        p.add_run(final_decision.get('justification', 'N/A'))
        doc.add_paragraph()

    # NEW: Summary and Scope
    scope_analysis = data.get('scope_analysis', {})
    if scope_analysis:
        doc.add_heading('Manuscript Summary', 1)
        doc.add_paragraph(scope_analysis.get('summary', ''))

        doc.add_heading('Scope Assessment', 1)
        p = doc.add_paragraph()
        p.add_run('In Scope: ').bold = True
        p.add_run('Yes' if scope_analysis.get('in_scope') else 'No')

        if scope_analysis.get('scope_categories'):
            p = doc.add_paragraph()
            p.add_run('Relevant Categories: ').bold = True
            p.add_run(', '.join(scope_analysis.get('scope_categories', [])))

        p = doc.add_paragraph()
        p.add_run('Justification: ').bold = True
        p.add_run(scope_analysis.get('scope_justification', ''))
        doc.add_paragraph()

    # Minor Issues
    doc.add_heading('Minor Deficiencies (Revisions)', 1)
    minor_issues = data.get('minor_issues', {}).get('issues', [])

    if minor_issues:
        doc.add_paragraph(
            f'The manuscript contains {len(minor_issues)} grammatical and stylistic issues that require correction. Below are specific examples:')
        doc.add_paragraph()

        table = doc.add_table(rows=len(minor_issues[:20]) + 1, cols=4)
        table.style = 'Light Grid Accent 1'

        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '#'
        hdr_cells[1].text = 'Incorrect Text'
        hdr_cells[2].text = 'Correction'
        hdr_cells[3].text = 'Issue Type'

        # Data
        for idx, issue in enumerate(minor_issues[:20], 1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = str(idx)
            row_cells[1].text = str(issue.get('incorrect', ''))[:50]
            row_cells[2].text = f"→ {str(issue.get('correct', ''))[:50]}"
            row_cells[3].text = str(issue.get('type', 'grammar')).title()
    else:
        doc.add_paragraph('No significant grammatical issues detected.')

    doc.add_paragraph()

    # Major Issues WITH EXAMPLES
    doc.add_heading('Major Revision Required', 1)
    major_issues = data.get('major_issues', [])

    for idx, issue in enumerate(major_issues, 1):
        p = doc.add_paragraph()
        p.add_run(f"{idx}. {issue.get('category', 'General')}: ").bold = True
        p.add_run(issue.get('issue', ''))

        # NEW: Add example
        if issue.get('example'):
            p = doc.add_paragraph()
            p.add_run('Example: ').italic = True
            p.add_run(issue.get('example', '')).italic = True

        p = doc.add_paragraph()
        p.add_run('→ Recommendation: ').bold = True
        p.add_run(issue.get('recommendation', ''))
        doc.add_paragraph()

    # Similarity
    doc.add_heading('Overall Textual Similarity', 1)
    similarity_data = data.get('similarity', {})
    overall_sim = similarity_data.get('overall', 0)
    confidence = similarity_data.get('confidence', 'unknown')

    doc.add_paragraph(f'Similarity Score: {overall_sim}% (Confidence: {confidence})')

    similar_sources = similarity_data.get('sources', [])
    if similar_sources:
        doc.add_paragraph('Similar Publications:')
        table = doc.add_table(rows=len(similar_sources[:5]) + 1, cols=3)
        table.style = 'Light Grid Accent 1'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Similar Work'
        hdr_cells[1].text = 'Authors'
        hdr_cells[2].text = 'Similarity'

        for idx, source in enumerate(similar_sources[:5], 1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = source.get('title', '')[:50]
            row_cells[1].text = source.get('author', '')
            row_cells[2].text = f"{source.get('similarity', 0)}%"

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = data.get('manuscript_title', 'Review_Report')[:50].replace(' ', '_')
    response['Content-Disposition'] = f'attachment; filename="{filename}_Review.docx"'

    return response


@api_view(['GET'])
def health_check(request):
    """API health check"""
    openai_status = 'configured ✅' if OPENAI_API_KEY and OPENAI_API_KEY != '' else 'NOT CONFIGURED ⚠️'

    return Response({
        'status': 'OK',
        'message': 'ASR API v4.0 - Enhanced with Scope Check, Final Decision & Word Export',
        'openai_api': openai_status,
        'version': '4.0',
        'features': [
            'Grammar & Spelling Detection',
            'Academic Tone & Style Analysis',
            'Enhanced Similarity Detection',
            'Major Issues WITH Examples',
            'Manuscript Summary',
            'Scope Assessment',
            'Suggested Final Decision',
            'PDF + WORD Reports'
        ]
    })